from __future__ import annotations

import re
import subprocess
import tempfile
from collections import Counter, defaultdict
from pathlib import Path

from pydantic import BaseModel, Field

from coalplan.application.writing_pattern_library import WritingPattern, WritingPatternLibrary, load_writing_pattern_library
from coalplan.domain.documents import stable_id


HEADING_RE = re.compile(
    r"^(?:第[一二三四五六七八九十百\d]+[章节篇]|"
    r"\d+(?:\.\d+)*[、.\s]+|"
    r"[一二三四五六七八九十]+[、.])\s*(?P<title>[\u4e00-\u9fffA-Za-z0-9（）()_—\-、/ ]{2,80})$"
)

PROJECT_TYPE_TERMS: dict[str, tuple[str, ...]] = {
    "煤火治理": ("煤火", "火区"),
    "市政管网": ("雨污", "管网", "市政", "道路", "排水"),
    "水电/抽蓄/隧洞边坡": ("水电", "抽水蓄能", "上水库", "导流", "交通洞", "隧洞", "边坡"),
    "新能源": ("光伏", "风电", "新能源"),
    "场平土石方": ("场平", "土石方", "PGA"),
    "爆破": ("爆破",),
}

PATTERN_TOPIC_TERMS: dict[str, tuple[str, ...]] = {
    "overview": ("工程概况", "工程概述", "项目概况", "建设规模", "主要工程量", "施工条件"),
    "deployment": ("施工部署", "施工组织", "组织机构", "总平面", "平面布置", "临建", "施工准备", "施工用水", "施工用电"),
    "craft": ("施工方法", "施工工艺", "施工程序", "主要施工", "钻孔", "灌浆", "注水", "混凝土", "支护", "开挖", "爆破", "管网", "安装", "调试"),
    "quality": ("质量", "检验", "试验", "检测", "验收"),
    "safety": ("安全", "职业健康", "危险源", "应急", "防洪", "度汛", "消防"),
    "environment": ("环保", "环境保护", "水土保持", "绿色施工", "文明施工", "扬尘", "噪声"),
    "schedule_resource": ("进度", "工期", "资源", "机械", "设备", "劳动力", "材料"),
}


BODY_CUE_TERMS: dict[str, tuple[tuple[str, tuple[str, ...]], ...]] = {
    "overview": (
        ("按项目背景、建设范围、主要工程量、现场条件组织概况正文", ("项目概况", "工程概况", "建设规模", "主要工程量", "施工条件")),
        ("把地形地貌、气象水文、交通条件写成施工条件约束", ("地形地貌", "水文气象", "交通条件", "施工场地", "地下管线")),
    ),
    "deployment": (
        ("按组织机构、职责分工、施工区段、总体顺序展开部署", ("组织机构", "职责", "施工区段", "施工部署", "总体安排")),
        ("把临建、道路、供水供电、通讯和材料堆场写为布置原则", ("施工总平面", "临时设施", "施工用水", "施工用电", "材料堆场")),
    ),
    "craft": (
        ("按施工准备、测量放样、工艺流程、过程控制、检查验收组织工艺正文", ("施工准备", "测量放样", "施工工艺", "工艺流程", "质量检查", "验收")),
        ("把人员、设备、材料、作业条件写入工序实施前置条件", ("人员配置", "机械设备", "主要材料", "作业条件", "资源配置")),
        ("把特殊情况处理、成品保护、记录资料作为工艺闭环", ("特殊情况处理", "成品保护", "施工记录", "资料整理")),
    ),
    "quality": (
        ("按质量目标、保证体系、责任分工、三检制、验收整改组织质量正文", ("质量目标", "质量保证体系", "责任制", "三检", "验收", "整改")),
        ("把原材料、工序、试验检测、隐蔽验收写成质量控制点", ("原材料", "工序质量", "试验检测", "隐蔽工程", "质量控制点")),
    ),
    "safety": (
        ("按安全目标、危险源、技术交底、现场检查、应急响应组织安全正文", ("安全目标", "危险源", "安全技术交底", "安全检查", "应急")),
        ("把消防、防汛、临电、机械、起重、爆破等专项风险单独成段", ("消防", "防汛", "临时用电", "机械伤害", "起重", "爆破")),
    ),
    "environment": (
        ("按扬尘、噪声、废水、固废、水土保持和文明施工分类组织措施", ("扬尘", "噪声", "废水", "固体废弃物", "水土保持", "文明施工")),
        ("把围挡、道路保洁、材料堆放、标识标牌写成现场标准化管理", ("围挡", "道路保洁", "材料堆放", "标识标牌", "标准化")),
    ),
    "schedule_resource": (
        ("按总体工期、阶段划分、关键线路、节点控制和纠偏措施组织进度正文", ("总工期", "阶段划分", "关键线路", "节点", "纠偏")),
        ("把劳动力、机械设备、材料供应和资金保障写成资源计划闭环", ("劳动力", "机械设备", "材料供应", "资金", "资源计划")),
    ),
}


class CorpusFileSummary(BaseModel):
    file_id: str
    file_name: str
    source_file: str | None = None
    extraction_status: str | None = None
    project_type: str
    heading_count: int
    headings: list[str] = Field(default_factory=list)
    matched_pattern_keys: list[str] = Field(default_factory=list)


class CorpusPatternStats(BaseModel):
    pattern_key: str
    file_count: int = 0
    heading_count: int = 0
    body_excerpt_count: int = 0
    common_headings: list[tuple[str, int]] = Field(default_factory=list)
    common_body_cues: list[tuple[str, int]] = Field(default_factory=list)
    source_files: list[str] = Field(default_factory=list)


class LocalCorpusAnalysis(BaseModel):
    corpus_dir: str
    corpus_content_kind: str = "toc"
    evidence_scope: str = "Local corpus is used for outline and writing-pattern guidance only; it is not a factual source for generated project content."
    sample_count: int
    project_type_counts: dict[str, int] = Field(default_factory=dict)
    high_frequency_headings: list[tuple[str, int]] = Field(default_factory=list)
    pattern_stats: dict[str, CorpusPatternStats] = Field(default_factory=dict)
    files: list[CorpusFileSummary] = Field(default_factory=list)


def analyze_local_corpus(
    corpus_dir: str | Path,
    *,
    include_source_excerpts: bool = False,
    max_source_chars: int = 250_000,
) -> LocalCorpusAnalysis:
    root = Path(corpus_dir)
    files = sorted(path for path in root.glob("*.txt") if path.name.lower() != "index.txt")
    file_summaries: list[CorpusFileSummary] = []
    project_type_counts: Counter[str] = Counter()
    all_headings: Counter[str] = Counter()
    pattern_heading_counts: dict[str, Counter[str]] = defaultdict(Counter)
    pattern_body_cue_counts: dict[str, Counter[str]] = defaultdict(Counter)
    pattern_body_file_counts: Counter[str] = Counter()
    pattern_files: dict[str, list[str]] = defaultdict(list)

    for path in files:
        text = path.read_text(encoding="utf-8", errors="ignore")
        metadata = extract_file_metadata(text)
        headings = extract_headings(text)
        normalized_headings = [normalize_heading(item) for item in headings]
        normalized_headings = [item for item in normalized_headings if item]
        project_type = classify_project_type(path.name)
        project_type_counts[project_type] += 1
        matched_patterns = sorted(match_patterns(normalized_headings))
        for heading in normalized_headings:
            all_headings[heading] += 1
        for key in matched_patterns:
            pattern_files[key].append(path.name)
            for heading in normalized_headings:
                if heading_matches_pattern(heading, key):
                    pattern_heading_counts[key][heading] += 1
            heading_cues = derive_body_cues_from_headings(normalized_headings, key)
            if heading_cues:
                pattern_body_file_counts[key] += 1
                for cue in heading_cues:
                    pattern_body_cue_counts[key][cue] += 1
        if include_source_excerpts and metadata.get("source_file"):
            source_text = extract_source_text(metadata["source_file"], max_chars=max_source_chars)
            if source_text:
                for key in matched_patterns:
                    cues = extract_body_cues(source_text, key)
                    if cues:
                        if not derive_body_cues_from_headings(normalized_headings, key):
                            pattern_body_file_counts[key] += 1
                        for cue in cues:
                            pattern_body_cue_counts[key][cue] += 1
        file_summaries.append(
            CorpusFileSummary(
                file_id=stable_id("corpusfile", path.name),
                file_name=path.name,
                source_file=metadata.get("source_file"),
                extraction_status=metadata.get("extraction_status"),
                project_type=project_type,
                heading_count=len(normalized_headings),
                headings=normalized_headings[:120],
                matched_pattern_keys=matched_patterns,
            )
        )

    pattern_stats = {
        key: CorpusPatternStats(
            pattern_key=key,
            file_count=len(set(pattern_files.get(key, []))),
            heading_count=sum(pattern_heading_counts.get(key, Counter()).values()),
            body_excerpt_count=pattern_body_file_counts.get(key, 0),
            common_headings=pattern_heading_counts.get(key, Counter()).most_common(20),
            common_body_cues=pattern_body_cue_counts.get(key, Counter()).most_common(20),
            source_files=sorted(set(pattern_files.get(key, [])))[:50],
        )
        for key in PATTERN_TOPIC_TERMS
    }
    return LocalCorpusAnalysis(
        corpus_dir=str(root),
        corpus_content_kind=_infer_corpus_content_kind(file_summaries),
        evidence_scope="Corpus samples are TOC/title extractions from local construction-organization documents. They train reusable outline coverage, subsection expansion, writing moves, required fact categories, and revision signals; generated project facts must still come from mapped input section_id/evidence_id or user supplements.",
        sample_count=len(files),
        project_type_counts=dict(project_type_counts.most_common()),
        high_frequency_headings=all_headings.most_common(80),
        pattern_stats=pattern_stats,
        files=file_summaries,
    )


def build_pattern_library_from_analysis(
    analysis: LocalCorpusAnalysis,
    *,
    base_library: WritingPatternLibrary | None = None,
) -> WritingPatternLibrary:
    base = base_library or load_writing_pattern_library()
    patterns: dict[str, WritingPattern] = {}
    for key, pattern in base.patterns.items():
        stats = analysis.pattern_stats.get(key)
        corpus_basis = list(pattern.corpus_basis)
        corpus_common_headings = list(pattern.corpus_common_headings)
        if stats:
            corpus_common_headings = [title for title, _count in stats.common_headings[:16]]
            corpus_basis = [
                f"本地语料样本数：{analysis.sample_count}",
                f"本模式命中施组文件：{stats.file_count}",
                f"正文片段可读样本：{stats.body_excerpt_count}",
                "高频相关标题：" + "；".join(f"{title}({count})" for title, count in stats.common_headings[:8])
                if stats.common_headings
                else "高频相关标题：未识别",
            ]
        auto_writable_moves = _merge_ordered(
            pattern.auto_writable_moves,
            _derive_auto_writable_moves(key, corpus_common_headings),
            _body_cues_to_auto_moves(key, [cue for cue, _count in stats.common_body_cues[:12]] if stats else []),
        )
        required_source_facts = _merge_ordered(
            pattern.required_source_facts,
            _derive_required_source_facts(key, corpus_common_headings),
        )
        revision_signals = _merge_ordered(
            pattern.revision_signals,
            _derive_revision_signals(key, corpus_common_headings),
            _body_cues_to_revision_signals(key, [cue for cue, _count in stats.common_body_cues[:12]] if stats else []),
        )
        patterns[key] = pattern.model_copy(
            update={
                "corpus_basis": corpus_basis,
                "corpus_common_headings": corpus_common_headings,
                "auto_writable_moves": auto_writable_moves,
                "required_source_facts": required_source_facts,
                "revision_signals": revision_signals,
            }
        )
    return WritingPatternLibrary(
        version=f"local-corpus-{analysis.sample_count}",
        corpus_scope=analysis.corpus_dir,
        patterns=patterns,
    )


def render_corpus_analysis_markdown(analysis: LocalCorpusAnalysis) -> str:
    lines = [
        "# Local Construction Organization Corpus Analysis",
        "",
        f"- corpus_dir: `{analysis.corpus_dir}`",
        f"- corpus_content_kind: {analysis.corpus_content_kind}",
        f"- evidence_scope: {analysis.evidence_scope}",
        f"- sample_count: {analysis.sample_count}",
        "",
        "## Project Types",
    ]
    for project_type, count in analysis.project_type_counts.items():
        lines.append(f"- {project_type}: {count}")
    lines.extend(["", "## Pattern Coverage"])
    for key, stats in analysis.pattern_stats.items():
        headings = "；".join(f"{title}({count})" for title, count in stats.common_headings[:8]) or "-"
        cues = "；".join(f"{cue}({count})" for cue, count in stats.common_body_cues[:6]) or "-"
        lines.append(
            f"- `{key}`: files={stats.file_count}, headings={stats.heading_count}, "
            f"body_excerpt_files={stats.body_excerpt_count}, common={headings}, body_cues={cues}"
        )
    lines.extend(["", "## Pattern Common Heading Seeds"])
    for key, stats in analysis.pattern_stats.items():
        headings = "；".join(title for title, _count in stats.common_headings[:12]) or "-"
        lines.append(f"- `{key}`: {headings}")
    lines.extend(["", "## Pattern Body Writing Cues"])
    for key, stats in analysis.pattern_stats.items():
        cues = "；".join(cue for cue, _count in stats.common_body_cues[:12]) or "-"
        lines.append(f"- `{key}`: {cues}")
    lines.extend(["", "## High Frequency Headings"])
    for heading, count in analysis.high_frequency_headings[:60]:
        lines.append(f"- {heading}: {count}")
    lines.extend(["", "## Files"])
    for item in analysis.files:
        patterns = ", ".join(item.matched_pattern_keys) or "-"
        source = f", source=`{item.source_file}`" if item.source_file else ""
        status = f", status={item.extraction_status}" if item.extraction_status else ""
        lines.append(f"- `{item.file_name}`: type={item.project_type}, headings={item.heading_count}, patterns={patterns}{status}{source}")
    return "\n".join(lines).strip() + "\n"


def extract_file_metadata(text: str) -> dict[str, str]:
    metadata: dict[str, str] = {}
    for raw in text.splitlines()[:20]:
        line = raw.strip()
        if _metadata_key_matches(line, "source_file"):
            metadata["source_file"] = _metadata_value(line)
        elif _metadata_key_matches(line, "extraction_status"):
            metadata["extraction_status"] = _metadata_value(line)
    return metadata


def extract_source_text(source_file: str | Path, *, max_chars: int = 250_000) -> str:
    path = Path(source_file)
    if not path.exists():
        return ""
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            return _extract_pdf_text(path, max_chars=max_chars)
        if suffix == ".docx":
            return _extract_docx_text(path, max_chars=max_chars)
    except Exception:
        return ""
    return ""


def extract_body_cues(text: str, pattern_key: str) -> list[str]:
    normalized = re.sub(r"\s+", "", text)
    cues: list[str] = []
    for cue, terms in BODY_CUE_TERMS.get(pattern_key, ()):
        if any(term in normalized for term in terms):
            cues.append(cue)
    return cues


def derive_body_cues_from_headings(headings: list[str], pattern_key: str) -> list[str]:
    """Infer reusable body-organization cues from TOC headings."""

    return extract_body_cues("\n".join(headings), pattern_key)


def _metadata_key_matches(line: str, key: str) -> bool:
    prefixes = {
        "source_file": ("源文件", "婧愭枃浠"),
        "extraction_status": ("提取状态", "鎻愬彇鐘舵"),
    }
    return any(line.startswith(prefix) for prefix in prefixes[key])


def _metadata_value(line: str) -> str:
    parts = re.split(r"[：:锛?]", line, maxsplit=1)
    return parts[1].strip() if len(parts) > 1 else ""


def _extract_pdf_text(path: Path, *, max_chars: int) -> str:
    with tempfile.TemporaryDirectory() as tmp:
        output = Path(tmp) / "out.txt"
        completed = subprocess.run(
            ["pdftotext", "-layout", "-enc", "UTF-8", str(path), str(output)],
            check=False,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="ignore",
            timeout=60,
        )
        if completed.returncode != 0 or not output.exists():
            return ""
        return output.read_text(encoding="utf-8", errors="ignore")[:max_chars]


def _extract_docx_text(path: Path, *, max_chars: int) -> str:
    try:
        from docx import Document
    except Exception:
        return ""
    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables[:20]:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)[:max_chars]


def extract_headings(text: str) -> list[str]:
    if "目录结构" in text:
        text = text.split("目录结构", 1)[1]
    headings: list[str] = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line or line.startswith(("源文件", "文件类型", "分类", "提取状态", "#")):
            continue
        line = re.sub(r"\s+\d+\s*$", "", line)
        line = line.strip(" ：:")
        if len(line) > 90:
            continue
        match = HEADING_RE.search(line)
        if match:
            headings.append(match.group("title").strip())
        elif _looks_like_plain_heading(line):
            headings.append(line)
    return headings


def normalize_heading(heading: str) -> str:
    text = re.sub(r"^第[一二三四五六七八九十百\d]+[章节篇]\s*", "", heading)
    text = re.sub(r"^\d+(?:\.\d+)*[、.\s]+", "", text)
    text = re.sub(r"^[一二三四五六七八九十]+[、.]", "", text)
    return text.strip(" ：:")


def classify_project_type(file_name: str) -> str:
    for project_type, terms in PROJECT_TYPE_TERMS.items():
        if any(term in file_name for term in terms):
            return project_type
    return "其他"


def match_patterns(headings: list[str]) -> set[str]:
    return {key for key in PATTERN_TOPIC_TERMS if any(heading_matches_pattern(heading, key) for heading in headings)}


def heading_matches_pattern(heading: str, pattern_key: str) -> bool:
    terms = PATTERN_TOPIC_TERMS.get(pattern_key, ())
    return any(term in heading for term in terms)


def _infer_corpus_content_kind(files: list[CorpusFileSummary]) -> str:
    if not files:
        return "unknown"
    statuses = " ".join(item.extraction_status or "" for item in files)
    if "目录" in statuses:
        return "toc_extraction"
    return "text_extraction"


def _derive_auto_writable_moves(pattern_key: str, headings: list[str]) -> list[str]:
    moves: list[str] = []
    if pattern_key == "overview":
        if _heading_has(headings, ("主要工程量", "工程数量", "工程量")):
            moves.append("把来源中的工程量、范围和施工对象整理成表格或条列。")
        if _heading_has(headings, ("施工条件", "地理位置", "交通", "水文", "气象")):
            moves.append("按自然条件、交通条件和现场条件组织工程概况段落。")
    elif pattern_key == "deployment":
        if _heading_has(headings, ("施工部署", "施工组织", "组织机构")):
            moves.append("按施工区段、组织机构、职责分工和作业顺序展开部署。")
        if _heading_has(headings, ("总平面", "临建", "施工用水", "施工用电")):
            moves.append("将总平面、临建、供水供电和道路条件整理为布置原则与人工待核项。")
    elif pattern_key == "craft":
        if _heading_has(headings, ("施工工艺流程", "施工程序", "施工工艺")):
            moves.append("按来源中的工艺流程和施工程序组织实施步骤。")
        if _heading_has(headings, ("施工设备", "人员配置", "材料", "资源")):
            moves.append("把人员、设备、材料和作业条件归并到施工准备与资源配置段落。")
        if _heading_has(headings, ("质量检查", "质量控制", "验收", "试验")):
            moves.append("把工序检查、试验检测和验收要求归并为质量控制闭环。")
        if _heading_has(headings, ("安全措施", "安全", "风险")):
            moves.append("将工艺相关风险转写为作业安全、设备安全和应急处置措施。")
    elif pattern_key == "quality":
        if _heading_has(headings, ("质量目标", "质量保证体系")):
            moves.append("按质量目标、组织体系、责任分工和制度流程展开。")
        if _heading_has(headings, ("质量控制", "质量检查", "验收", "试验")):
            moves.append("按材料进场、工序控制、试验检测、验收整改组织质量措施。")
    elif pattern_key == "safety":
        if _heading_has(headings, ("安全目标", "安全保证体系", "安全管理")):
            moves.append("按安全目标、组织体系、岗位职责和检查整改闭环展开。")
        if _heading_has(headings, ("应急预案", "消防", "防汛", "度汛", "危险源")):
            moves.append("按风险识别、预警响应、应急资源、处置流程和演练记录组织正文。")
    elif pattern_key == "environment":
        if _heading_has(headings, ("文明施工", "环境保护", "绿色施工")):
            moves.append("按扬尘、噪声、废水、固废、生态保护和现场标准化分类展开措施。")
        if _heading_has(headings, ("水土保持", "弃渣", "弃土")):
            moves.append("将弃渣弃土、水土保持和地方监管要求写为人工核验优先的控制项。")
    elif pattern_key == "schedule_resource":
        if _heading_has(headings, ("施工进度计划", "工期目标", "工期保证")):
            moves.append("按总体工期、关键线路、节点控制和纠偏措施组织进度正文。")
        if _heading_has(headings, ("资源配置", "机械", "设备", "劳动力", "材料")):
            moves.append("把机械设备、劳动力和材料计划整理为资源保障与进场控制。")
    return moves


def _derive_required_source_facts(pattern_key: str, headings: list[str]) -> list[str]:
    facts: list[str] = []
    if _heading_has(headings, ("主要工程量", "工程数量", "工程量")):
        facts.extend(["工程量", "单位"])
    if _heading_has(headings, ("施工工艺流程", "施工程序", "施工工艺")):
        facts.extend(["工艺流程", "施工顺序"])
    if _heading_has(headings, ("施工设备", "机械", "设备")):
        facts.append("设备配置")
    if _heading_has(headings, ("劳动力", "人员配置")):
        facts.append("劳动力配置")
    if _heading_has(headings, ("质量检查", "质量控制", "验收", "试验")):
        facts.extend(["检验项目", "验收要求"])
    if _heading_has(headings, ("安全", "危险源", "应急", "消防", "防汛", "度汛")):
        facts.extend(["危险源", "安全措施"])
    if _heading_has(headings, ("环境保护", "文明施工", "水土保持", "弃渣", "弃土")):
        facts.extend(["环保控制对象", "文明施工要求"])
    if pattern_key == "schedule_resource" and _heading_has(headings, ("进度", "工期", "节点")):
        facts.extend(["工期", "节点"])
    if pattern_key == "overview" and _heading_has(headings, ("地理位置", "工程概况", "项目概况")):
        facts.extend(["项目位置", "施工范围"])
    return facts


def _derive_revision_signals(pattern_key: str, headings: list[str]) -> list[str]:
    signals: list[str] = []
    if pattern_key == "craft" and _heading_has(headings, ("施工工艺流程", "施工程序", "质量检查")):
        signals.append("正文未按工艺流程、质量检查和验收闭环展开。")
    if pattern_key == "quality" and _heading_has(headings, ("质量目标", "质量保证体系", "质量控制")):
        signals.append("质量章节缺少目标、体系、过程控制或检查整改闭环。")
    if pattern_key == "safety" and _heading_has(headings, ("危险源", "应急预案", "消防", "防汛")):
        signals.append("安全章节缺少危险源识别、应急响应或专项风险措施。")
    if pattern_key == "schedule_resource" and _heading_has(headings, ("施工进度计划", "资源配置")):
        signals.append("进度资源章节缺少节点安排、资源配置或纠偏保障措施。")
    return signals


def _body_cues_to_auto_moves(pattern_key: str, cues: list[str]) -> list[str]:
    return [f"参考原始施组正文样本：{cue}" for cue in cues]


def _body_cues_to_revision_signals(pattern_key: str, cues: list[str]) -> list[str]:
    return [f"正文未体现本类施组常见写法：{cue}" for cue in cues]


def _heading_has(headings: list[str], terms: tuple[str, ...]) -> bool:
    return any(any(term in heading for term in terms) for heading in headings)


def _merge_ordered(*groups: list[str]) -> list[str]:
    merged: list[str] = []
    for group in groups:
        for item in group:
            if item and item not in merged:
                merged.append(item)
    return merged


def _looks_like_plain_heading(line: str) -> bool:
    if len(line) < 3 or len(line) > 40:
        return False
    if not re.search(r"[\u4e00-\u9fff]", line):
        return False
    noise = ("。", "，", "；", "、如下", "见下", "公司", "有限公司")
    return not any(term in line for term in noise)
